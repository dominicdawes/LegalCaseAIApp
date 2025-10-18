"""For the new 'True In-Memory Streaming' method"""

# utils/document_loaders/docx_loader.py

import textract
import os
import io
import tempfile
from typing import Iterator, Union, List
from langchain.schema import Document
from .base import BaseDocumentLoader
import docx # python-docx
import docx2txt
from celery.utils.log import get_task_logger

# ——— Logging & Env Load ———————————————————————————————————————————————————————————
logger = get_task_logger(__name__)
logger.propagate = False

# 📏 This is your "magic number" heuristic.
# You will need to tune this based on your typical documents.
CHARS_PER_PAGE_ESTIMATE = 3000

# ——— Loader Classes ———————————————————————————————————————————————————————————

class DocxLoader(BaseDocumentLoader):
    """
    ### MODIFIED: A loader for .docx files that operates on in-memory streams.
    This loader does NOT support the old binary .doc format for in-memory processing.
    """
    def stream_documents(self, source: Union[str, io.BytesIO]) -> Iterator[Document]:
        """
        ### CHANGE: Uses docx.Document() with a file-like object.
        """
        try:
            document = docx.Document(source)
            
            # 1. Add a tracker for cumulative character count
            cumulative_chars = 0
            
            for i, para in enumerate(document.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                
                # 2. Update character count
                # Add 2 for an assumed newline/paragraph break
                cumulative_chars += len(text) + 2
                
                # 3. Calculate the estimated page number
                estimated_page = (cumulative_chars // CHARS_PER_PAGE_ESTIMATE) + 1
                logger.debug(f"DOC/DOCX estimated_pages: {estimated_page}")
                
                yield Document(
                    page_content=text,
                    metadata={
                        "paragraph_index": i,
                        # 4. Add the new metadata field
                        "estimated_page": estimated_page,
                        "page": estimated_page
                    },
                )
        except Exception as e:
            raise RuntimeError(f"Failed to process DOCX stream: {e}")

class LegacyDocLoader(BaseDocumentLoader):
    """A loader for legacy .doc files that "converts"??? from .doc to .docx and then
    uses Textract for in-memory streming."""
    def stream_documents(self, source: io.BytesIO) -> Iterator[Document]:
        # Add some basic validation
        if not isinstance(source, io.BytesIO):
            raise ValueError("LegacyDocLoader requires BytesIO input")
            
        source.seek(0)
        if len(source.getvalue()) == 0:
            raise ValueError("Empty document provided")
        
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            source.seek(0)
            tmp.write(source.read())
            tmp.flush()
            
            try:
                raw_bytes = textract.process(tmp.name)
                text = raw_bytes.decode('utf-8', errors='ignore')
                
                # Slightly more robust paragraph splitting
                paragraphs = [p.strip() for p in text.split('\n\n') if p.strip() and len(p.strip()) > 10]
                
                logger.info(f"Extracted {len(paragraphs)} paragraphs from LEGACY .doc file")
                
            except Exception as e:
                raise RuntimeError(f"Failed to process legacy .doc file: {e}")
            finally:
                try:
                    os.unlink(tmp.name)
                except OSError:
                    pass  # File might already be gone
            
            # Stream the results
            for i, para in enumerate(paragraphs):
                yield Document(
                    page_content=para, 
                    metadata={
                        "paragraph_index": i, 
                        "source_type": "legacy_doc",
                        "total_paragraphs": len(paragraphs)
                    }
                )